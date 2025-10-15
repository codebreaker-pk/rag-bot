from __future__ import annotations
from pathlib import Path
from io import BytesIO
from typing import Dict, List, Tuple

from rag import chunk, add_docs

SUPPORTED = {".pdf", ".docx", ".txt"}

def _read_text_from_file(fp: Path, max_pdf_pages: int = 20) -> str:
    """Safe readers; never raise to caller."""
    ext = fp.suffix.lower()
    try:
        if ext == ".pdf":
            from pypdf import PdfReader
            text = ""
            with fp.open("rb") as f:
                reader = PdfReader(f)
                for p in reader.pages[:max_pdf_pages]:
                    try:
                        text += (p.extract_text() or "") + "\n"
                    except Exception:
                        pass
            return text
        elif ext == ".docx":
            from docx import Document
            doc = Document(str(fp))
            return "\n".join(par.text for par in doc.paragraphs)
        elif ext == ".txt":
            return fp.read_text(encoding="utf-8", errors="ignore")
        else:
            return ""
    except Exception:
        return ""

def extract_text_from_bytes(filename: str, data: bytes, max_pdf_pages: int = 20) -> str:
    ext = (filename or "").lower()
    try:
        if ext.endswith(".pdf"):
            from pypdf import PdfReader
            text = ""
            reader = PdfReader(BytesIO(data))
            for p in reader.pages[:max_pdf_pages]:
                try:
                    text += (p.extract_text() or "") + "\n"
                except Exception:
                    pass
            return text
        elif ext.endswith(".docx"):
            from docx import Document
            doc = Document(BytesIO(data))
            return "\n".join(par.text for par in doc.paragraphs)
        elif ext.endswith(".txt"):
            return data.decode("utf-8", errors="ignore")
        else:
            return ""
    except Exception:
        return ""

def load_and_ingest(domain: str, dirpath: str) -> Dict:
    """
    Ingest all supported files under dirpath into vector store, in small batches.
    Returns stats dict (never raises).
    """
    d = Path(dirpath)
    stats = {"domain": domain, "files_indexed": 0, "chunks_indexed": 0, "errors": []}
    if not d.exists():
        return stats

    files = [p for p in d.rglob("*") if p.is_file() and p.suffix.lower() in SUPPORTED]

    for fp in files:
        text = _read_text_from_file(fp)
        if not text.strip():
            stats["errors"].append({"file": fp.name, "error": "no text extracted"})
            continue

        pieces = chunk(text)
        if not pieces:
            stats["errors"].append({"file": fp.name, "error": "no chunks"})
            continue

        docs = [{
            "id": f"{domain}:{fp.name}:{i}",
            "title": fp.name,
            "source": f"{domain}/{fp.name}",
            "text": piece
        } for i, piece in enumerate(pieces)]

        BATCH = 64
        for i in range(0, len(docs), BATCH):
            batch = docs[i:i+BATCH]
            try:
                add_docs(domain, batch)
                stats["chunks_indexed"] += len(batch)
            except Exception as e:
                stats["errors"].append({"file": fp.name, "error": f"add_docs failed: {e}"})
                break

        stats["files_indexed"] += 1

    return stats
